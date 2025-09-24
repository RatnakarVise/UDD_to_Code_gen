from docx import Document
from docx.shared import Pt, RGBColor
from io import BytesIO
import re

# ---------------------------------------------------
# Style Helpers
# ---------------------------------------------------

def add_heading(doc, text):
    paragraph = doc.add_paragraph()
    run = paragraph.add_run(text)
    run.bold = True
    run.underline = True
    run.font.color.rgb = RGBColor(0, 0, 255)
    run.font.size = Pt(14)

def add_subheading(doc, text):
    paragraph = doc.add_paragraph()
    run = paragraph.add_run(text)
    run.bold = True
    run.font.size = Pt(12)

def add_paragraph(doc, text):
    paragraph = doc.add_paragraph()
    cursor = 0
    for match in re.finditer(r"\*\*(.+?)\*\*", text):
        start, end = match.span()
        paragraph.add_run(text[cursor:start])
        bold_run = paragraph.add_run(match.group(1))
        bold_run.bold = True
        cursor = end
    paragraph.add_run(text[cursor:])

def add_code_block(doc, code_lines):
    para = doc.add_paragraph()
    run = para.add_run("\n".join(code_lines))
    run.font.name = "Courier New"
    run.font.size = Pt(10)

def add_markdown_table(doc, lines):
    headers = [cell.strip(" *") for cell in lines[0].split("|") if cell.strip()]
    rows = [
        [cell.strip() for cell in row.split("|") if cell.strip()]
        for row in lines[2:]
    ]
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
    for row in rows:
        row_cells = table.add_row().cells
        for i, cell in enumerate(row):
            row_cells[i].text = cell

# ---------------------------------------------------
# Main Parser/Generator
# ---------------------------------------------------

def generate_doc(ts_text: str, file_obj: BytesIO, title: str):
    doc = Document()
    doc.add_heading(title, level=1)

    lines = ts_text.splitlines()
    current_section = ""
    current_content = []
    seen_sections = set()

    in_code_block = False
    code_block_lines = []
    in_table = False
    table_lines = []

    section_header_pattern = re.compile(r"^(\d+)\.\s+(.+)$")         # e.g., 3. Business Requirement
    subheading_pattern = re.compile(r"^(\d+)\.(\d+)\s+(.+)$")        # e.g., 3.1 Requirement Overview
    table_line_pattern = re.compile(r"^\|(.+?)\|$")

    known_main_sections = set()

    def flush_section():
        nonlocal current_section
        if current_section and current_section not in seen_sections:
            add_heading(doc, current_section)
            seen_sections.add(current_section)
        for para in current_content:
            add_paragraph(doc, para)
        current_content.clear()

    for line in lines:
        line = line.strip()
        if not line:
            continue

        # Handle code blocks
        if isinstance(line, bytes):
            line = line.decode("utf-8")
        if line.startswith("```"):
            in_code_block = not in_code_block
            if not in_code_block:
                add_code_block(doc, code_block_lines)
                code_block_lines.clear()
            continue
        elif in_code_block:
            code_block_lines.append(line)
            continue

        # Handle tables
        if table_line_pattern.match(line):
            table_lines.append(line)
            in_table = True
            continue
        elif in_table:
            flush_section()
            add_markdown_table(doc, table_lines)
            table_lines.clear()
            in_table = False
            continue

        # Main section (e.g., 3. Business Requirement)
        main_match = section_header_pattern.match(line)
        if main_match:
            flush_section()
            current_section = f"{main_match.group(1)}. {main_match.group(2)}"
            known_main_sections.add(main_match.group(1))
            continue

        # Subsection (e.g., 3.1 Requirement Overview)
        sub_match = subheading_pattern.match(line)
        if sub_match:
            flush_section()
            main_number = sub_match.group(1)
            sub_number = sub_match.group(2)
            sub_title = sub_match.group(3)

            if main_number not in known_main_sections:
                # Insert missing main section heading only once
                main_title = f"{main_number}. (Auto) Section"
                if main_title not in seen_sections:
                    add_heading(doc, main_title)
                    seen_sections.add(main_title)
                    known_main_sections.add(main_number)

            subheading_text = f"{main_number}.{sub_number} {sub_title}"
            add_subheading(doc, subheading_text)
            continue

        # Regular content
        current_content.append(line)

    # Final flush
    flush_section()
    doc.save(file_obj)



# ----------------------------------------
# Public API Functions for FS, TS, ABAP
# ----------------------------------------

def create_functional_spec_docx(fs_text: str, file_obj: BytesIO):
    generate_doc(fs_text, file_obj, title="FUNCTIONAL SPECIFICATION")

def create_technical_spec_docx(ts_text: str, file_obj: BytesIO):
    generate_doc(ts_text, file_obj, title="TECHNICAL SPECIFICATION")

def create_abap_code_docx(abap_code: str, file_obj: BytesIO):
    """
    Creates an ABAP code DOCX file in the given file object, using monospaced font.
    """
    doc = Document()
    doc.add_heading('ABAP Code', level=1)
    para = doc.add_paragraph()
    run = para.add_run(abap_code)
    run.font.name = 'Consolas'  # Monospaced font for code
    run.font.size = Pt(10.5)
    doc.save(file_obj)
