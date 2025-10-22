from docx import Document
from docx.shared import Pt

def create_abap_code_docx(abap_code: str, file_obj):
    """
    Creates an ABAP code DOCX file in the given file object.
    Accepts a single string containing all ABAP code.
    """
    doc = Document()
    doc.add_heading('ABAP Code', level=1)

    # Add the whole ABAP code as a single monospaced block
    para = doc.add_paragraph()
    run = para.add_run(abap_code)
    run.font.name = 'Consolas'
    run.font.size = Pt(10.5)

    doc.save(file_obj)
